"""
查詢結果匯出模組

支援匯出格式：
- Markdown (.md)
- JSON (.json)
- PDF (.pdf)
- Word (.docx)
- Excel (.xlsx)
"""

import json
from datetime import datetime
from io import BytesIO


class ReportExporter:
    """查詢結果匯出器"""

    def __init__(self, data: dict):
        """
        初始化匯出器

        Args:
            data: 查詢結果資料（包含 regulations、query 等）
        """
        self.data = data
        self.query = data.get("query", "未知查詢")
        self.original_query = data.get("original_query", self.query.split("\n")[0])
        self.timestamp = data.get("timestamp", datetime.now().isoformat())

        # 取得法規資料
        regulations = data.get("regulations", {})
        if isinstance(regulations, dict):
            self.verified_regulations = regulations.get("verified_regulations", [])
            self.timeline = regulations.get("timeline", [])
            self.compliance_checklist = regulations.get("compliance_checklist", [])
            self.summary = regulations.get("summary", "")
            self.confidence_score = regulations.get("confidence_score", 0)
        else:
            self.verified_regulations = regulations if isinstance(regulations, list) else []
            self.timeline = []
            self.compliance_checklist = []
            self.summary = ""
            self.confidence_score = 0

    def to_markdown(self) -> str:
        """
        匯出為 Markdown 格式

        Returns:
            Markdown 字串
        """
        lines = []

        # 標題
        lines.append("# 法規查詢報告\n")
        lines.append(f"**查詢內容**: {self.original_query}\n")
        lines.append(f"**查詢時間**: {self.timestamp}\n")

        # 摘要
        if self.summary:
            lines.append("\n## 摘要\n")
            lines.append(f"{self.summary}\n")

        # 相關法規
        if self.verified_regulations:
            lines.append(f"\n## 相關法規 ({len(self.verified_regulations)} 項)\n")
            for i, reg in enumerate(self.verified_regulations, 1):
                name = reg.get("name", "未知")
                name_zh = reg.get("name_zh", "")
                url = reg.get("url", "")
                reg_type = reg.get("type", "")

                lines.append(f"### {i}. {name}\n")
                if name_zh and name_zh != name:
                    lines.append(f"**中文名稱**: {name_zh}\n")
                if reg_type:
                    lines.append(f"**類型**: {reg_type}\n")
                if url:
                    lines.append(f"**來源**: {url}\n")

                # 重點
                key_points = reg.get("key_points", [])
                if key_points:
                    lines.append("\n**重點摘要**:\n")
                    for point in key_points:
                        lines.append(f"- {point}\n")

                # 條文節錄
                excerpts = reg.get("article_excerpts", [])
                if excerpts:
                    lines.append("\n**條文節錄**:\n")
                    for excerpt in excerpts:
                        article_num = excerpt.get("article_number", "")
                        content = excerpt.get("content", "")
                        if article_num:
                            lines.append(f"\n**{article_num}**\n")
                        if content:
                            lines.append(f"> {content}\n")

                lines.append("\n---\n")

        # 時間軸
        if self.timeline:
            lines.append("\n## 法規時間軸\n")
            lines.append("| 日期 | 事件 | 相關法規 |\n")
            lines.append("|------|------|----------|\n")
            for event in self.timeline:
                date = event.get("date", "未知")
                event_desc = event.get("event", "")
                regulation = event.get("regulation", "")
                lines.append(f"| {date} | {event_desc} | {regulation} |\n")

        # 合規檢核清單
        if self.compliance_checklist:
            lines.append("\n## 合規檢核清單\n")
            for i, item in enumerate(self.compliance_checklist, 1):
                item_name = item.get("item", "")
                description = item.get("description", "")
                priority = item.get("priority", "medium")
                priority_icon = {"high": "[!]", "medium": "[*]", "low": "[-]"}.get(priority, "[*]")

                lines.append(f"{i}. {priority_icon} **{item_name}**\n")
                if description:
                    lines.append(f"   - {description}\n")

        # 信心分數
        if self.confidence_score:
            lines.append("\n---\n")
            lines.append(f"*分析信心度: {int(self.confidence_score * 100)}%*\n")

        return "".join(lines)

    def to_json(self, indent: int = 2) -> str:
        """
        匯出為 JSON 格式

        Args:
            indent: 縮排空格數

        Returns:
            JSON 字串
        """
        export_data = {
            "export_info": {
                "format": "json",
                "exported_at": datetime.now().isoformat(),
                "query": self.original_query,
            },
            "result": self.data
        }
        return json.dumps(export_data, ensure_ascii=False, indent=indent)

    def to_pdf(self) -> bytes:
        """
        匯出為 PDF 格式

        Returns:
            PDF 檔案的 bytes
        """
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import cm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

        # 嘗試註冊中文字體
        chinese_font = 'Helvetica'  # 預設值
        font_paths = [
            # macOS 系統字體（優先使用 STHeiti，較穩定）
            ('/System/Library/Fonts/STHeiti Light.ttc', 0),
            ('/System/Library/Fonts/STHeiti Medium.ttc', 0),
            # PingFang（subfontIndex=0 為簡體，1 為繁體）
            ('/System/Library/Fonts/PingFang.ttc', 1),
            # Arial Unicode（單一 TTF）
            ('/Library/Fonts/Arial Unicode.ttf', None),
            # Linux 系統字體
            ('/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc', 0),
            ('/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf', None),
        ]

        for font_path, subfont_index in font_paths:
            try:
                import os
                if os.path.exists(font_path):
                    if subfont_index is not None:
                        pdfmetrics.registerFont(TTFont('Chinese', font_path, subfontIndex=subfont_index))
                    else:
                        pdfmetrics.registerFont(TTFont('Chinese', font_path))
                    chinese_font = 'Chinese'
                    break
            except Exception:
                continue

        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=2*cm, bottomMargin=2*cm)

        styles = getSampleStyleSheet()

        # 自訂樣式
        title_style = ParagraphStyle(
            'ChineseTitle',
            parent=styles['Title'],
            fontName=chinese_font,
            fontSize=18,
            spaceAfter=20,
        )
        heading_style = ParagraphStyle(
            'ChineseHeading',
            parent=styles['Heading2'],
            fontName=chinese_font,
            fontSize=14,
            spaceBefore=15,
            spaceAfter=10,
        )
        body_style = ParagraphStyle(
            'ChineseBody',
            parent=styles['Normal'],
            fontName=chinese_font,
            fontSize=10,
            leading=14,
        )

        story = []

        # 標題
        story.append(Paragraph("法規查詢報告", title_style))
        story.append(Paragraph(f"查詢: {self.original_query}", body_style))
        story.append(Paragraph(f"時間: {self.timestamp}", body_style))
        story.append(Spacer(1, 20))

        # 摘要
        if self.summary:
            story.append(Paragraph("摘要", heading_style))
            story.append(Paragraph(self.summary, body_style))
            story.append(Spacer(1, 10))

        # 相關法規
        if self.verified_regulations:
            story.append(Paragraph(f"相關法規 ({len(self.verified_regulations)} 項)", heading_style))

            for i, reg in enumerate(self.verified_regulations, 1):
                name = reg.get("name", "未知")
                name_zh = reg.get("name_zh", "")
                url = reg.get("url", "")

                reg_text = f"<b>{i}. {name}</b>"
                if name_zh and name_zh != name:
                    reg_text += f"<br/>{name_zh}"
                if url:
                    reg_text += f"<br/>來源: {url}"

                story.append(Paragraph(reg_text, body_style))

                # 重點
                key_points = reg.get("key_points", [])
                if key_points:
                    for point in key_points[:3]:  # 限制數量
                        story.append(Paragraph(f"  - {point}", body_style))

                story.append(Spacer(1, 10))

        # 合規檢核清單
        if self.compliance_checklist:
            story.append(Paragraph("合規檢核清單", heading_style))

            # 使用 Paragraph 讓文字自動換行
            table_cell_style = ParagraphStyle(
                'TableCell',
                parent=styles['Normal'],
                fontName=chinese_font,
                fontSize=8,
                leading=10,
                wordWrap='CJK',  # 支援中文換行
            )

            table_data = [[
                Paragraph("<b>項目</b>", table_cell_style),
                Paragraph("<b>說明</b>", table_cell_style),
                Paragraph("<b>優先級</b>", table_cell_style),
            ]]
            for item in self.compliance_checklist[:8]:  # 限制數量
                item_text = item.get("item", "")
                desc_text = item.get("description", "")
                priority = item.get("priority", "medium")

                table_data.append([
                    Paragraph(item_text, table_cell_style),
                    Paragraph(desc_text, table_cell_style),
                    Paragraph(priority, table_cell_style),
                ])

            # A4 可用寬度約 17cm，設定總寬度為 16cm
            table = Table(table_data, colWidths=[4.5*cm, 9.5*cm, 2*cm])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('FONTNAME', (0, 0), (-1, -1), chinese_font),
                ('FONTSIZE', (0, 0), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                ('TOPPADDING', (0, 0), (-1, -1), 4),
                ('BOTTOMPADDING', (0, 1), (-1, -1), 4),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
            ]))
            story.append(table)

        # 信心分數
        if self.confidence_score:
            story.append(Spacer(1, 20))
            story.append(Paragraph(f"分析信心度: {int(self.confidence_score * 100)}%", body_style))

        doc.build(story)
        return buffer.getvalue()

    def to_docx(self) -> bytes:
        """
        匯出為 Word 格式

        Returns:
            DOCX 檔案的 bytes
        """
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # 標題
        title = doc.add_heading("法規查詢報告", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 基本資訊
        doc.add_paragraph(f"查詢內容: {self.original_query}")
        doc.add_paragraph(f"查詢時間: {self.timestamp}")

        # 摘要
        if self.summary:
            doc.add_heading("摘要", level=1)
            doc.add_paragraph(self.summary)

        # 相關法規
        if self.verified_regulations:
            doc.add_heading(f"相關法規 ({len(self.verified_regulations)} 項)", level=1)

            for i, reg in enumerate(self.verified_regulations, 1):
                name = reg.get("name", "未知")
                name_zh = reg.get("name_zh", "")
                url = reg.get("url", "")
                reg_type = reg.get("type", "")

                doc.add_heading(f"{i}. {name}", level=2)

                if name_zh and name_zh != name:
                    doc.add_paragraph(f"中文名稱: {name_zh}")
                if reg_type:
                    doc.add_paragraph(f"類型: {reg_type}")
                if url:
                    doc.add_paragraph(f"來源: {url}")

                # 重點
                key_points = reg.get("key_points", [])
                if key_points:
                    doc.add_paragraph("重點摘要:")
                    for point in key_points:
                        doc.add_paragraph(f"  - {point}")

                # 條文節錄
                excerpts = reg.get("article_excerpts", [])
                if excerpts:
                    doc.add_paragraph("條文節錄:")
                    for excerpt in excerpts:
                        article_num = excerpt.get("article_number", "")
                        content = excerpt.get("content", "")
                        if article_num:
                            p = doc.add_paragraph()
                            p.add_run(article_num).bold = True
                        if content:
                            doc.add_paragraph(content, style='Quote')

        # 時間軸
        if self.timeline:
            doc.add_heading("法規時間軸", level=1)
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '日期'
            hdr_cells[1].text = '事件'
            hdr_cells[2].text = '相關法規'

            for event in self.timeline:
                row_cells = table.add_row().cells
                row_cells[0].text = event.get("date", "")
                row_cells[1].text = event.get("event", "")
                row_cells[2].text = event.get("regulation", "")

        # 合規檢核清單
        if self.compliance_checklist:
            doc.add_heading("合規檢核清單", level=1)
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '項目'
            hdr_cells[1].text = '說明'
            hdr_cells[2].text = '優先級'

            for item in self.compliance_checklist:
                row_cells = table.add_row().cells
                row_cells[0].text = item.get("item", "")
                row_cells[1].text = item.get("description", "")
                row_cells[2].text = item.get("priority", "")

        # 信心分數
        if self.confidence_score:
            doc.add_paragraph()
            doc.add_paragraph(f"分析信心度: {int(self.confidence_score * 100)}%")

        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def to_xlsx(self) -> bytes:
        """
        匯出為 Excel 格式

        Returns:
            XLSX 檔案的 bytes
        """
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill

        wb = Workbook()

        # Sheet 1: 概覽
        ws_overview = wb.active
        ws_overview.title = "概覽"

        ws_overview['A1'] = "法規查詢報告"
        ws_overview['A1'].font = Font(size=16, bold=True)
        ws_overview['A3'] = "查詢內容"
        ws_overview['B3'] = self.original_query
        ws_overview['A4'] = "查詢時間"
        ws_overview['B4'] = self.timestamp
        ws_overview['A5'] = "信心分數"
        ws_overview['B5'] = f"{int(self.confidence_score * 100)}%" if self.confidence_score else "N/A"

        if self.summary:
            ws_overview['A7'] = "摘要"
            ws_overview['A7'].font = Font(bold=True)
            ws_overview['A8'] = self.summary

        ws_overview.column_dimensions['A'].width = 15
        ws_overview.column_dimensions['B'].width = 60

        # Sheet 2: 法規列表
        if self.verified_regulations:
            ws_regs = wb.create_sheet("法規列表")

            headers = ["編號", "名稱", "中文名稱", "類型", "來源", "重點摘要"]
            header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")

            for col, header in enumerate(headers, 1):
                cell = ws_regs.cell(row=1, column=col, value=header)
                cell.font = Font(color="FFFFFF", bold=True)
                cell.fill = header_fill

            for row, reg in enumerate(self.verified_regulations, 2):
                ws_regs.cell(row=row, column=1, value=row-1)
                ws_regs.cell(row=row, column=2, value=reg.get("name", ""))
                ws_regs.cell(row=row, column=3, value=reg.get("name_zh", ""))
                ws_regs.cell(row=row, column=4, value=reg.get("type", ""))
                ws_regs.cell(row=row, column=5, value=reg.get("url", ""))

                key_points = reg.get("key_points", [])
                ws_regs.cell(row=row, column=6, value="\n".join(key_points) if key_points else "")

            # 調整欄寬
            ws_regs.column_dimensions['A'].width = 8
            ws_regs.column_dimensions['B'].width = 40
            ws_regs.column_dimensions['C'].width = 30
            ws_regs.column_dimensions['D'].width = 15
            ws_regs.column_dimensions['E'].width = 50
            ws_regs.column_dimensions['F'].width = 50

        # Sheet 3: 合規檢核清單
        if self.compliance_checklist:
            ws_checklist = wb.create_sheet("合規檢核清單")

            headers = ["編號", "項目", "說明", "優先級", "法規依據"]
            header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")

            for col, header in enumerate(headers, 1):
                cell = ws_checklist.cell(row=1, column=col, value=header)
                cell.font = Font(color="FFFFFF", bold=True)
                cell.fill = header_fill

            for row, item in enumerate(self.compliance_checklist, 2):
                ws_checklist.cell(row=row, column=1, value=row-1)
                ws_checklist.cell(row=row, column=2, value=item.get("item", ""))
                ws_checklist.cell(row=row, column=3, value=item.get("description", ""))
                ws_checklist.cell(row=row, column=4, value=item.get("priority", ""))
                ws_checklist.cell(row=row, column=5, value=item.get("regulation_basis", ""))

            ws_checklist.column_dimensions['A'].width = 8
            ws_checklist.column_dimensions['B'].width = 30
            ws_checklist.column_dimensions['C'].width = 50
            ws_checklist.column_dimensions['D'].width = 10
            ws_checklist.column_dimensions['E'].width = 30

        # Sheet 4: 時間軸
        if self.timeline:
            ws_timeline = wb.create_sheet("時間軸")

            headers = ["日期", "事件", "相關法規"]
            header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")

            for col, header in enumerate(headers, 1):
                cell = ws_timeline.cell(row=1, column=col, value=header)
                cell.font = Font(color="FFFFFF", bold=True)
                cell.fill = header_fill

            for row, event in enumerate(self.timeline, 2):
                ws_timeline.cell(row=row, column=1, value=event.get("date", ""))
                ws_timeline.cell(row=row, column=2, value=event.get("event", ""))
                ws_timeline.cell(row=row, column=3, value=event.get("regulation", ""))

            ws_timeline.column_dimensions['A'].width = 15
            ws_timeline.column_dimensions['B'].width = 50
            ws_timeline.column_dimensions['C'].width = 30

        buffer = BytesIO()
        wb.save(buffer)
        return buffer.getvalue()


def export_result(data: dict, format: str) -> tuple[bytes | str, str, str]:
    """
    便捷函數：匯出查詢結果

    Args:
        data: 查詢結果資料
        format: 匯出格式 (markdown, json, pdf, docx, xlsx)

    Returns:
        (檔案內容, 檔案名稱, MIME 類型)
    """
    exporter = ReportExporter(data)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    query_short = data.get("original_query", data.get("query", "report"))[:20].replace(" ", "_")

    if format == "markdown":
        content = exporter.to_markdown()
        filename = f"法規報告_{query_short}_{timestamp}.md"
        mime = "text/markdown"
    elif format == "json":
        content = exporter.to_json()
        filename = f"法規報告_{query_short}_{timestamp}.json"
        mime = "application/json"
    elif format == "pdf":
        content = exporter.to_pdf()
        filename = f"法規報告_{query_short}_{timestamp}.pdf"
        mime = "application/pdf"
    elif format == "docx":
        content = exporter.to_docx()
        filename = f"法規報告_{query_short}_{timestamp}.docx"
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif format == "xlsx":
        content = exporter.to_xlsx()
        filename = f"法規報告_{query_short}_{timestamp}.xlsx"
        mime = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    else:
        raise ValueError(f"不支援的格式: {format}")

    return content, filename, mime
